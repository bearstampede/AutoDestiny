#! python3
import os, sys, time, inspect, datetime
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog, QApplication, QWidget, QPushButton, QMessageBox
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import pandas as pd
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches
from win32com.client import DispatchEx

excelPrettyReportMacro = """
Sub Macro2()
'
 
'
    Application.Calculation = xlCalculationManual
    'Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False

    ActiveSheet.Cells.UnMerge
    Range("F5").Select
    Selection.Cut Destination:=Range("V7")
    Rows("1:6").Select
    Selection.Delete Shift:=xlUp
    Columns("A:A").Select
    Selection.Delete Shift:=xlToLeft
    Range("T1").Select
    Selection.Copy
    Range("T1").Select
    Application.CutCopyMode = False
    Selection.Cut
    Range("S1").Select
    ActiveSheet.Paste
    ActiveWindow.SmallScroll Down:=-3
    Columns("A:A").Select
    Selection.NumberFormat = "[$-F800]dddd, mmmm dd, yyyy"
    Selection.ColumnWidth = 29
'date changes'
'Edit & cut out excess colummes '

    Columns("D:F").Select
    Application.ReplaceFormat.Clear
    Selection.Replace What:="am", Replacement:=" am", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=False
    Selection.Replace What:="pm", Replacement:=" pm", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=False
    Selection.NumberFormat = "[$-409]h:mm AM/PM;@"
    Columns("O:O").Select
    Selection.Cut Destination:=Columns("B:B")
    Columns("B:B").Select
    Selection.Replace What:="Classroom", Replacement:="", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=False
    Selection.ColumnWidth = 20.86
    Selection.Replace What:="Conference Room", Replacement:="CR ", LookAt:= _
        xlPart, SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=False
    Range("C:C,E:E").Select
    Range("E1").Activate
    Selection.Delete Shift:=xlToLeft
    Range("C12").Select
    Columns("B:B").ColumnWidth = 20.71
    Columns("B:B").ColumnWidth = 7
    Columns("I:I").Select
    Selection.Cut Destination:=Columns("E:E")
    Columns("J:J").Select
    Selection.Cut Destination:=Columns("F:F")
    Columns("F:F").Select
    Columns("G:G").EntireColumn.AutoFit
    Columns("O:O").Select
    Selection.Cut Destination:=Columns("H:H")
    Columns("Q:Q").Select
    Selection.Cut Destination:=Columns("I:I")
    Columns("N:N").Select
    Selection.Cut Destination:=Columns("J:J")
    Columns("P:P").Select
    Columns("P:P").Cut Destination:=Columns("K:K")
    Columns("L:R").Select
    Selection.Delete Shift:=xlToLeft
    Columns("A:L").Select
    Selection.ColumnWidth = 12.57
    Columns("G:G").Select
    Selection.ColumnWidth = 16
    Columns("H:H").Select
    Selection.ColumnWidth = 18.5
    Range("B:B,F:F").Select
    Range("F1").Activate
    With Selection.Font
        .Color = -16777024
        .TintAndShade = 0
    End With
    Range("C1").Select
    ActiveCell.FormulaR1C1 = "Start"
    Range("D1").Select
    ActiveCell.FormulaR1C1 = "End"
    Range("D2").Select
    Columns("C:C").ColumnWidth = 6.71
    Columns("C:D").Select
    Columns("C:D").EntireColumn.AutoFit
    Range("A21").Select
    Selection.End(xlDown).Select
    ActiveWindow.SmallScroll Down:=3
    Columns("C:C").Select
    Selection.Insert Shift:=xlToRight, CopyOrigin:=xlFormatFromLeftOrAbove
    Columns("L:L").Select
    Selection.Cut Destination:=Columns("C:C")
    Range("C1").Select
    ActiveCell.FormulaR1C1 = "Size"
    Range("C2").Select
    Columns("C:C").EntireColumn.AutoFit
    Range("B6").Select

'
    Columns("A:A").ColumnWidth = 29
    Columns("B:B").ColumnWidth = 7.29
    Columns("D:E").Select
    With Application.ReplaceFormat.Font
        .Subscript = False
        .Color = 255
        .TintAndShade = 0
    End With
    Selection.Replace What:="AM", Replacement:="AM", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Range("H20").Select


' Sort Fields DATES and then Room numbers '

    Dim LASTROW As Long
    LASTROW = Cells(Rows.Count, 2).End(xlUp).Row
    Cells.Select
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Clear
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("A2:A" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("D2:D" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("B2:B" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal
    With ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort
        .SetRange Range("A1:AA" & LASTROW)
        .Header = xlYes
        .MatchCase = False
        .Orientation = xlTopToBottom
        .SortMethod = xlPinYin
        .Apply
    End With

'
    Columns("B:B").Select
    Selection.Font.Bold = False
    Selection.Font.Bold = True
    Columns("D:E").Select
    With Application.ReplaceFormat.Interior
        .PatternColorIndex = xlAutomatic
        .ThemeColor = xlThemeColorAccent5
        .TintAndShade = 0.799981688894314
        .PatternTintAndShade = 0
    End With
    Selection.Replace What:="AM", Replacement:="AM", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Columns("B:B").Select
    With Application.ReplaceFormat.Interior
        .PatternColorIndex = xlAutomatic
        .Color = 65535
        .TintAndShade = 0
        .PatternTintAndShade = 0
    End With
    Selection.Replace What:="201", Replacement:="201", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Selection.Replace What:="502", Replacement:="502", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Selection.Replace What:="510", Replacement:="510", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Selection.Replace What:="514", Replacement:="514", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Selection.Replace What:="515", Replacement:="515", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Range("F25").Select


'
    Application.DisplayAlerts = False
    Cells.Select
    Selection.Subtotal GroupBy:=1, Function:=xlCount, TotalList:=Array(4), _
        Replace:=False, PageBreaks:=False, SummaryBelowData:=False
    Range("G26").Select
    Application.DisplayAlerts = True

'
    Rows("2:2").Select
    Selection.Delete Shift:=xlUp
    Columns("A:A").Select
    Application.ReplaceFormat.Clear
    With Application.ReplaceFormat.Interior
        .PatternColorIndex = xlAutomatic
        .Color = 5287936
        .TintAndShade = 0
        .PatternTintAndShade = 0
    End With
    Selection.Replace What:="Count", Replacement:=" -- ", LookAt:=xlPart, _
        SearchOrder:=xlByRows, MatchCase:=False, SearchFormat:=False, _
        ReplaceFormat:=True
    Range("E12").Select
    Columns("F:F").ColumnWidth = 57
    Columns("F:F").ColumnWidth = 46.29
    Columns("I:K").Select
    With Selection
        .WrapText = True
        .Orientation = 0
        .AddIndent = False
        .IndentLevel = 0
        .ShrinkToFit = False
        .ReadingOrder = xlContext
        .MergeCells = False
    End With
    With Selection
        .WrapText = False
        .Orientation = 0
        .AddIndent = False
        .IndentLevel = 0
        .ShrinkToFit = False
        .ReadingOrder = xlContext
        .MergeCells = False
    End With
    Rows("1:1").Select
    Selection.Insert Shift:=xlDown, CopyOrigin:=xlFormatFromLeftOrAbove
    Range("M2").Select
    Selection.Cut Destination:=Range("A1")
    Range("A1").Select
    Selection.Font.Bold = True
    Range("D24").Select


' Freeze Panes top 2 rows '

    Range("A1:D1").Select
    With Selection.Interior
        .Pattern = xlSolid
        .PatternColorIndex = xlAutomatic
        .Color = 49407
        .TintAndShade = 0
        .PatternTintAndShade = 0
    End With
    Range("A2").Select
    ActiveWindow.FreezePanes = True
    ActiveWindow.SmallScroll Down:=-3
    Range("A3").Select
    ActiveWindow.FreezePanes = False
    ActiveWindow.FreezePanes = True
    ActiveWindow.SmallScroll Down:=0
    Range("A3").Select


' set date for Report generated as of  '

    Range("F1").Select
    ActiveCell.FormulaR1C1 = "Report generated as of   "
    Range("F1").Select
    With Selection
        .HorizontalAlignment = xlRight
        .VerticalAlignment = xlBottom
        .WrapText = False
        .Orientation = 0
        .AddIndent = False
        .IndentLevel = 0
        .ShrinkToFit = False
        .ReadingOrder = xlContext
        .MergeCells = False
    End With
    Range("G1").Select
    ActiveCell.FormulaR1C1 = "=R[2]C[-6]"
    Range("F1:I1").Select
    With Selection.Interior
        .Pattern = xlSolid
        .PatternColorIndex = xlAutomatic
        .Color = 5287936
        .TintAndShade = 0
        .PatternTintAndShade = 0
    End With
    Selection.Font.Bold = True
    Range("C12").Select
    ActiveWindow.SmallScroll Down:=0
    Range("F1:I1").Select
    With Selection.Interior
        .Pattern = xlSolid
        .PatternColorIndex = xlAutomatic
        .ThemeColor = xlThemeColorAccent6
        .TintAndShade = 0.799981688894314
        .PatternTintAndShade = 0
    End With
    Range("F1:I1").Select
    Selection.Copy
    Selection.PasteSpecial Paste:=xlPasteValues, Operation:=xlNone, SkipBlanks _
        :=False, Transpose:=False
    ActiveSheet.Paste
    Application.CutCopyMode = False
    
    Range(Cells(1, 12), Cells(1, Columns.Count)).Columns.Hidden = True
    
    Columns("D:E").Cells.HorizontalAlignment = xlHAlignRight
    Columns("K:K").ColumnWidth = 56
    
    Columns("A:A").Find("Destiny One CLOUD Production report generated by ").Select
    Rows(ActiveCell.Row & ":" & Rows.Count).Delete Shift:=xlUp
    Rows(ActiveCell.Row + 1 & ":" & Rows.Count).EntireRow.Hidden = True
    ActiveSheet.Cells(3, 1).Select
    
    Application.EnableEvents = True
    Application.DisplayStatusBar = True
    'Application.ScreenUpdating = True
    Application.Calculation = xlCalculationAutomatic

End Sub
"""

excelClassScheduleMacro = """

Sub ReformatDestinyReport()
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False

    Call TurnOffMerge
    
    Dim dateVar
    'dateVar = Range("F3").Text
    'dateVar = Format(dateVar, "m/dd/yyyy")

    ' Delete top n rows.  Find B column with "Date" string.  The location of this row depends on parameters supplied to report.
    Dim SearchRange As Range
    Dim FindRow As Range
    Set SearchRange = Range("B1", Range("B65536").End(xlUp))
    Set FindRow = SearchRange.Find("Date", LookIn:=xlValues, lookat:=xlWhole)
    Rows("1:" + Trim(Str(FindRow.Row) - 1)).Delete
    
    ' Delete last row
    Set rlastrow = Cells(Rows.Count, "A").End(xlUp)
    Cells.Find(What:="*", After:=[A1], SearchOrder:=xlByRows, SearchDirection:=xlPrevious).EntireRow.Delete
    
    ' Delete blank columns
    Call DeleteBlankColumns
    
    ' Freeze top header cells
    ' http://stackoverflow.com/questions/3232920/how-can-i-programmatically-freeze-the-top-row-of-an-excel-worksheet-in-excel-200
    'Application.ScreenUpdating = True
    Dim r As Range
    'Set r = ActiveCell
    Rows("2:2").Select
    With ActiveWindow
        .FreezePanes = False
        .ScrollRow = 1
        .ScrollColumn = 1
        .FreezePanes = True
        '.ScrollRow = r.Row
    End With
    'r.Select
    
    ' Delete unneeded columns
    Range("O1").EntireColumn.Delete
    Range("N1").EntireColumn.Delete
    Range("M1").EntireColumn.Delete
    Range("L1").EntireColumn.Delete
    Range("K1").EntireColumn.Delete
    Range("J1").EntireColumn.Delete
    Range("H1").EntireColumn.Delete
    Range("B1").EntireColumn.Delete
    'Range("A1").EntireColumn.Delete
    
    ' Make time columns general format rather than text
    Call ConvertToGeneral
    
    ' Make time columns appear in time format rather than text so they can be sorted
    Call ConvertTextColumnsToTime

    Dim LASTROW As Long
    LASTROW = Cells(Rows.Count, 2).End(xlUp).Row
    'Cells.Select
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Clear
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("A2:A" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("B2:B" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal
    ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort.SortFields.Add Key:=Range("G2:G" & LASTROW _
        ), SortOn:=xlSortOnValues, Order:=xlAscending, DataOption:=xlSortNormal

    With ActiveWorkbook.Worksheets(ActiveSheet.Name).Sort
        .SetRange Range("A1:AA" & LASTROW)
        .Header = xlYes
        .MatchCase = False
        .Orientation = xlTopToBottom
        .SortMethod = xlPinYin
        .Apply
    End With

    Columns("F").Replace What:="Instructor To Be Announced", Replacement:="TBA", LookAt:=xlPart, SearchOrder:=xlByRows, MatchCase:=False

    Dim i As Long
    Dim prevDate
    Dim currentSheet As Worksheet, ws As Worksheet
    Dim multiDate As Boolean
    Dim currentSheetName, newDate As String
    multiDate = False

    currentSheetName = Format(Range("A2").Value, "m-dd-yyyy")
    ActiveSheet.Name = currentSheetName
    Set currentSheet = ActiveWorkbook.Sheets(currentSheetName)
    prevDate = currentSheet.Range("A2").Value
    For i = 3 To Range("A" & Rows.Count).End(xlUp).Row
        If currentSheet.Range("A" & i).Value <> prevDate Then
            multiDate = True
            prevDate = currentSheet.Range("A" & i).Value
            newDate = Format(currentSheet.Range("A" & i).Value, "m-dd-yyyy")
            ActiveWorkbook.Sheets.Add After:=Worksheets(Worksheets.Count)
            ActiveWorkbook.Sheets(Worksheets.Count).Name = newDate
            Set ws = Worksheets(newDate)
            'Range("A2").Select
            Rows("2:2").Select
            With ActiveWindow
                .FreezePanes = False
                .ScrollRow = 1
                .ScrollColumn = 1
                .FreezePanes = True
                '.ScrollRow = r.Row
            End With
            currentSheet.Cells(1, 1).EntireRow.Copy ws.Range("A1")
        End If
        If multiDate = True Then
            currentSheet.Range("A" & i).EntireRow.Cut ws.Range("A" & ws.Cells(ws.Rows.Count, "A").End(xlUp).Row + 1)
        End If
    Next i

    For Each ws In Worksheets
        With ws
            dateVar = Format(ws.Range("A2").Value, "dddd mmmm d, yyyy")
            .Range("A1").EntireColumn.Delete
            Call addRows(ws, dateVar)

            '.Columns("D:D").ColumnWidth = 58
            ' Increase font size to 18 for all cells
            With .Cells.Font
                .Name = "Consolas"
                .Size = "18"
            End With
            
            Const xlLandscape = 2

            Dim LR As Long
            LR = .Range("A" & Rows.Count).End(xlUp).Row
            With .PageSetup
                .CenterHorizontally = True
                .CenterVertically = True
                .PrintArea = "A1:F" & LR
                .PaperSize = xlPaperLetter
                .Orientation = xlLandscape
                .Zoom = False
                .LeftMargin = Application.InchesToPoints(0.25)
                .RightMargin = Application.InchesToPoints(0.25)
                .TopMargin = Application.InchesToPoints(0.25)
                .BottomMargin = Application.InchesToPoints(0.25)
                .HeaderMargin = Application.InchesToPoints(0)
                .FooterMargin = Application.InchesToPoints(0)
                .PrintGridlines = True
                .FitToPagesWide = 1
                .FitToPagesTall = 1
            End With

            ' Set columns to autofit
                .Columns("A:B").NumberFormat = "h:mm AM/PM"
                .Columns("A:A").WrapText = False
                .Columns("D:D").WrapText = True
                .Columns("B:C").Columns.AutoFit
                'Columns("F:F").ColumnWidth = 35
                .Columns("E:F").Columns.AutoFit
                .Range("A2:A2").Columns.AutoFit
                .Columns("D:D").ColumnWidth = 62
                '.Range("A1:A" & LR).EntireRow.AutoFit
                '.Columns("D:D").ColumnWidth = 60
                '.Columns("D:D").ColumnWidth = .Columns("D:D").ColumnWidth + 2
                '.Columns("D:D").Columns.AutoFit
                .Range("A1:F" & LR).Rows.AutoFit
        End With
    Next ws

    Application.EnableEvents = True
    Application.DisplayStatusBar = True
    Application.ScreenUpdating = True
    Application.Calculation = xlCalculationAutomatic

    ' Activate top cell
    Sheets(1).Activate
  
End Sub

Sub TurnOffMerge()
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False
    Dim bProtected As Boolean
    On Error GoTo TryActiveProtectedViewWindow
    If ActiveWorkbook.Excel8CompatibilityMode Then
      bProtected = ActiveSheet.ProtectContents
      GoTo Continue
    End If
    
TryActiveProtectedViewWindow:
    On Error GoTo GiveUp
    bProtected = Application.ActiveProtectedViewWindow <> nil
    GoTo Continue

GiveUp:
    MsgBox "There was an error using this function."
    On Error Resume Next
    Exit Sub
    
    
    

    'If Application.ActiveProtectedViewWindow <> nil Then
Continue:
    If bProtected Then
      MsgBox "You must turn off protected view before you can use this function"
    Else


        Cells.MergeCells = False
        For Each mCell In ActiveSheet.UsedRange.Rows(1).Cells
            mCell.EntireColumn.AutoFit
        Next mCell

    End If

End Sub

Sub DeleteBlankColumns()
    Dim Col As Long, ColCnt As Long, Rng As Range
     
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False
     
On Error GoTo Exits:
     
    If Selection.Columns.Count > 1 Then
        Set Rng = Selection
    Else
        Set Rng = Range(Columns(1), Columns(ActiveSheet.Cells.SpecialCells(xlCellTypeLastCell).Column()))
    End If
    ColCnt = 0
    For Col = Rng.Columns.Count To 1 Step -1
        If Application.WorksheetFunction.CountA(Rng.Columns(Col).EntireColumn) = 0 Then
            Rng.Columns(Col).EntireColumn.Delete
            ColCnt = ColCnt + 1
        End If
    Next Col
     
Exits:
    Application.EnableEvents = True
    Application.DisplayStatusBar = True
    Application.ScreenUpdating = True
    Application.Calculation = xlCalculationAutomatic

End Sub

Sub ConvertToGeneral()
' converttogeneral Macro
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False

    ActiveWindow.SmallScroll Down:=-6
    Range("B2:B" & Cells(Rows.Count, "B").End(xlUp).Row).NumberFormat = "General"
    
    Range("C2:C" & Cells(Rows.Count, "C").End(xlUp).Row).NumberFormat = "General"

End Sub

Sub ConvertTextColumnsToTime()
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False

    'For Each c In Range("B2:C" & Cells(Rows.Count, "A").End(xlUp).Row)
    For Each c In Range("B2:B" & Cells(Rows.Count, "B").End(xlUp).Row)
        c.Value = TimeValue(c.Value)
    Next
    'For Each c In Range("B2:B" & Cells(Rows.Count, "B").End(xlUp).Row)
    For Each c In Range("C2:C" & Cells(Rows.Count, "C").End(xlUp).Row)
        c.Value = TimeValue(c.Value)
    Next
    'Application.ScreenUpdating = False
End Sub

Sub addRows(ws As Worksheet, dateVar)
    Application.Calculation = xlCalculationManual
    Application.ScreenUpdating = False
    Application.DisplayStatusBar = False
    Application.EnableEvents = False
    Dim i As Long
    'Set ws

    'Application.ScreenUpdating = False

    ' Loop through rows in reverse, starting from the end
    For i = ws.Range("A" & Rows.Count).End(xlUp).Row To 2 Step -1
        Select Case True
            ' Evening Classes
            Case ws.Range("A" & i - 1).Value < TimeValue("5:45 PM") And ws.Range("A" & i).Value >= TimeValue("5:45 PM")
                ws.Rows(i).Resize(1).Insert
                ws.Range("A" & i).NumberFormat = "General"
                ws.Range("A" & i).Font.Bold = True
                ws.Range("A" & i).Value = "Evening Classes"

            ' Afternoon Classes
            Case ws.Range("A" & i - 1).Value < TimeValue("12:00 PM") And ws.Range("A" & i).Value >= TimeValue("12:00 PM")
                ws.Rows(i).Resize(1).Insert
                ws.Range("A" & i).NumberFormat = "General"
                ws.Range("A" & i).Font.Bold = True
                ws.Range("A" & i).Value = "Afternoon Classes"

            ' Morning Classes
            Case i = 2 And ws.Range("A" & i).Value >= TimeValue("8:00 AM")
                ws.Rows(i).Resize(1).Insert
                ws.Range("A" & i).NumberFormat = "General"
                ws.Range("A" & i).Font.Bold = True
                ws.Range("A" & i).Value = "Morning Classes"
        End Select
    Next

    ' Add row at the top to add date
    ws.Rows(1).Resize(2).Insert
    ws.Range("A1").NumberFormat = "General"
    ws.Range("A1").Font.Bold = True
    ws.Range("A1").Value = "UC Berkeley Extension  " & dateVar & Chr(10)
    'Application.ScreenUpdating = True
End Sub

"""

current_folder = os.path.realpath(os.path.abspath(os.path.split(inspect.getfile(inspect.currentframe() ))[0]))
ABSWTemplate = os.path.join(current_folder, "Template-ABSW.docx")
GBCTemplate = os.path.join(current_folder, "Template-GBC.docx")
genReport = False
startDate = '2018-01-01'
endDate = '2018-01-01'
location = 'ABSW'
createSigns = False
useExistingReport = False
saveReportToPath = ''
existingReportPath, _ = ('','')
singleOutput = False
saveSignsDirectory = ''
center = {'ABSW': {'campus': 'Berkeley - CA0001', 'building': 'UC Berkeley Extension American Baptist Seminary of the West, 2515 Hillegass Ave. - '},
          'Belmont': {'campus': 'Belmont - CA0004', 'building': 'UC Berkeley Extension Belmont Center, 1301 Shoreway Rd., Ste. 400 - BEL'},
          'Golden Bear Center': {'campus': 'Berkeley - CA0001', 'building': 'UC Berkeley Extension Golden Bear Center, 1995 University Ave. - GBC'},
          'San Francisco Center': {'campus': 'San Francisco - CA0003', 'building': 'San Francisco Campus, 160 Spear St. - SFCAMPUS'}
            }
##centerReverse = {'ABSW - UC Berkeley Extension American Baptist Seminary of the West, 2515 Hillegass Ave.': 'ABSW',
##                 'BEL - UC Berkeley Extension Belmont Center, 1301 Shoreway Rd., Ste. 400': 'Belmont',
##                 'GBC - UC Berkeley Extension Golden Bear Center, 1995 University Ave.': 'Golden Bear Center',
##                 'SFCAMPUS - San Francisco Campus, 160 Spear St.': 'San Francisco Center'}

centerReverse = {'ABSW - UC Berkeley Extension American Baptist Seminary of the West, 2515 Hillegass Ave.': {'name':'ABSW', 'template': ABSWTemplate},
                 'BEL - UC Berkeley Extension Belmont Center, 1301 Shoreway Rd., Ste. 400': {'name':'BLM', 'template': ABSWTemplate},
                 'GBC - UC Berkeley Extension Golden Bear Center, 1995 University Ave.': {'name': 'GBC', 'template' : GBCTemplate},
                 'SFCAMPUS - San Francisco Campus, 160 Spear St.': {'name': 'SFC', 'template': ABSWTemplate}}

# Main Window for GUI
class Ui_mainWindow(object): 
    def setupUi(self, mainWindow):
        global startDate
        global endDate
        mainWindow.setObjectName("mainWindow")
        mainWindow.setWindowModality(QtCore.Qt.NonModal)
        mainWindow.setEnabled(True)
        mainWindow.resize(430, 380)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(mainWindow.sizePolicy().hasHeightForWidth())
        mainWindow.setSizePolicy(sizePolicy)
        mainWindow.setMinimumSize(QtCore.QSize(430, 380))
        mainWindow.setMaximumSize(QtCore.QSize(430, 380))
        mainWindow.setBaseSize(QtCore.QSize(430, 325))
        self.gridLayoutWidget = QtWidgets.QWidget(mainWindow)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(10, 10, 412, 371))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.widgetVerticalLayout = QtWidgets.QVBoxLayout(self.gridLayoutWidget)
        self.widgetVerticalLayout.setSizeConstraint(QtWidgets.QLayout.SetDefaultConstraint)
        self.widgetVerticalLayout.setContentsMargins(0, 0, 0, 0)
        self.widgetVerticalLayout.setObjectName("widgetVerticalLayout")
        self.genReportBox = QtWidgets.QGroupBox(self.gridLayoutWidget)
        self.genReportBox.setMinimumSize(QtCore.QSize(410, 120))
        self.genReportBox.setMaximumSize(QtCore.QSize(410, 120))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(14)
        font.setUnderline(True)
        self.genReportBox.setFont(font)
        self.genReportBox.setCheckable(True)
        self.genReportBox.setChecked(False)
        self.genReportBox.setObjectName("genReportBox")
        self.selectLocation = QtWidgets.QComboBox(self.genReportBox)
        self.selectLocation.setGeometry(QtCore.QRect(90, 60, 171, 22))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectLocation.setFont(font)
        self.selectLocation.setEditable(False)
        self.selectLocation.setObjectName("selectLocation")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.startDateLabel = QtWidgets.QLabel(self.genReportBox)
        self.startDateLabel.setGeometry(QtCore.QRect(20, 30, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.startDateLabel.setFont(font)
        self.startDateLabel.setObjectName("startDateLabel")
        self.browseSaveReportButton = QtWidgets.QToolButton(self.genReportBox)
        self.browseSaveReportButton.setGeometry(QtCore.QRect(340, 90, 61, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setUnderline(False)
        self.browseSaveReportButton.setFont(font)
        self.browseSaveReportButton.setCheckable(False)
        self.browseSaveReportButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseSaveReportButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseSaveReportButton.setObjectName("browseSaveReportButton")
        self.selectSaveReportPath = QtWidgets.QLineEdit(self.genReportBox)
        self.selectSaveReportPath.setGeometry(QtCore.QRect(100, 90, 241, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectSaveReportPath.setFont(font)
        self.selectSaveReportPath.setReadOnly(False)
        self.selectSaveReportPath.setObjectName("selectSaveReportPath")
        self.selectStartDate = QtWidgets.QDateEdit(self.genReportBox)
        self.selectStartDate.setGeometry(QtCore.QRect(100, 30, 111, 22))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setBold(False)
        font.setUnderline(False)
        font.setWeight(50)
        self.selectStartDate.setFont(font)
        self.selectStartDate.setFrame(True)
        self.selectStartDate.setProperty("showGroupSeparator", False)
        self.selectStartDate.setCalendarPopup(True)
        self.selectStartDate.setDate(QtCore.QDate.currentDate())
        startDate = str(QtCore.QDate.currentDate().toPyDate())
        self.selectStartDate.setObjectName("selectStartDate")
        self.locationLabel = QtWidgets.QLabel(self.genReportBox)
        self.locationLabel.setGeometry(QtCore.QRect(20, 60, 71, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.locationLabel.setFont(font)
        self.locationLabel.setObjectName("locationLabel")
        self.saveReportPathLabel = QtWidgets.QLabel(self.genReportBox)
        self.saveReportPathLabel.setGeometry(QtCore.QRect(20, 90, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.saveReportPathLabel.setFont(font)
        self.saveReportPathLabel.setObjectName("saveReportPathLabel")
        self.endDateLabel = QtWidgets.QLabel(self.genReportBox)
        self.endDateLabel.setGeometry(QtCore.QRect(220, 30, 71, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setBold(False)
        font.setUnderline(False)
        font.setWeight(50)
        self.endDateLabel.setFont(font)
        self.endDateLabel.setObjectName("endDateLabel")
        self.selectEndDate = QtWidgets.QDateEdit(self.genReportBox)
        self.selectEndDate.setGeometry(QtCore.QRect(290, 30, 110, 22))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectEndDate.setFont(font)
        self.selectEndDate.setFocusPolicy(QtCore.Qt.WheelFocus)
        self.selectEndDate.setCalendarPopup(True)
        self.selectEndDate.setDate(QtCore.QDate.currentDate())
        endDate = str(QtCore.QDate.currentDate().toPyDate())
        self.selectEndDate.setObjectName("selectEndDate")
        self.selectLocation.raise_()
        self.startDateLabel.raise_()
        self.browseSaveReportButton.raise_()
        self.selectSaveReportPath.raise_()
        self.selectStartDate.raise_()
        self.locationLabel.raise_()
        self.saveReportPathLabel.raise_()
        self.endDateLabel.raise_()
        self.selectEndDate.raise_()
        self.widgetVerticalLayout.addWidget(self.genReportBox)
        self.createSignsBox = QtWidgets.QGroupBox(self.gridLayoutWidget)
        self.createSignsBox.setEnabled(True)
        self.createSignsBox.setMinimumSize(QtCore.QSize(150, 160))
        self.createSignsBox.setMaximumSize(QtCore.QSize(410, 150))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(14)
        font.setUnderline(True)
        self.createSignsBox.setFont(font)
        self.createSignsBox.setCheckable(True)
        self.createSignsBox.setChecked(False)
        self.createSignsBox.setObjectName("createSignsBox")
        self.singleOutputCheckbox = QtWidgets.QCheckBox(self.createSignsBox)
        self.singleOutputCheckbox.setGeometry(QtCore.QRect(20, 100, 151, 21))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.singleOutputCheckbox.setFont(font)
        self.singleOutputCheckbox.setChecked(False)
        self.singleOutputCheckbox.setObjectName("singleOutputCheckbox")
        self.useExistingReportBox = QtWidgets.QGroupBox(self.createSignsBox)
        self.useExistingReportBox.setGeometry(QtCore.QRect(10, 30, 390, 60))
        self.useExistingReportBox.setMaximumSize(QtCore.QSize(390, 60))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.useExistingReportBox.setFont(font)
        self.useExistingReportBox.setCheckable(True)
        self.useExistingReportBox.setChecked(False)
        self.useExistingReportBox.setObjectName("useExistingReportBox")
        self.selectExistingReportPath = QtWidgets.QLineEdit(self.useExistingReportBox)
        self.selectExistingReportPath.setGeometry(QtCore.QRect(10, 30, 311, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.selectExistingReportPath.setFont(font)
        self.selectExistingReportPath.setReadOnly(False)
        self.selectExistingReportPath.setObjectName("selectExistingReportPath")
        self.browseExistingReportButton = QtWidgets.QToolButton(self.useExistingReportBox)
        self.browseExistingReportButton.setGeometry(QtCore.QRect(320, 30, 61, 19))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.browseExistingReportButton.setFont(font)
        self.browseExistingReportButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseExistingReportButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseExistingReportButton.setAutoRaise(False)
        self.browseExistingReportButton.setArrowType(QtCore.Qt.NoArrow)
        self.browseExistingReportButton.setObjectName("browseExistingReportButton")
        self.selectExistingReportPath.raise_()
        self.browseExistingReportButton.raise_()
        self.genReportBox.raise_()
        self.genReportBox.raise_()
        self.selectSaveSignsPath = QtWidgets.QLineEdit(self.createSignsBox)
        self.selectSaveSignsPath.setGeometry(QtCore.QRect(100, 130, 231, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectSaveSignsPath.setFont(font)
        self.selectSaveSignsPath.setReadOnly(False)
        self.selectSaveSignsPath.setObjectName("selectSaveSignsPath")
        self.browseSaveSignsButton = QtWidgets.QToolButton(self.createSignsBox)
        self.browseSaveSignsButton.setGeometry(QtCore.QRect(330, 130, 61, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setUnderline(False)
        self.browseSaveSignsButton.setFont(font)
        self.browseSaveSignsButton.setCheckable(False)
        self.browseSaveSignsButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseSaveSignsButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseSaveSignsButton.setObjectName("browseSaveSignsButton")
        self.saveSignsPathLabel = QtWidgets.QLabel(self.createSignsBox)
        self.saveSignsPathLabel.setGeometry(QtCore.QRect(20, 130, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.saveSignsPathLabel.setFont(font)
        self.saveSignsPathLabel.setObjectName("saveSignsPathLabel")
        self.widgetVerticalLayout.addWidget(self.createSignsBox)
        self.startExitLayout = QtWidgets.QGridLayout()
        self.startExitLayout.setObjectName("startExitLayout")
        self.exitButton = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.exitButton.setMinimumSize(QtCore.QSize(0, 50))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        self.exitButton.setFont(font)
        self.exitButton.setObjectName("exitButton")
        self.startExitLayout.addWidget(self.exitButton, 0, 1, 1, 1)
        self.StartButton = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.StartButton.setMinimumSize(QtCore.QSize(0, 50))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        self.StartButton.setFont(font)
        self.StartButton.setObjectName("StartButton")
        self.startExitLayout.addWidget(self.StartButton, 0, 0, 1, 1)
        self.widgetVerticalLayout.addLayout(self.startExitLayout)

        self.retranslateUi(mainWindow)
        QtCore.QMetaObject.connectSlotsByName(mainWindow)

    def retranslateUi(self, mainWindow):
        _translate = QtCore.QCoreApplication.translate
        mainWindow.setWindowTitle(_translate("mainWindow", "Auto Schedule"))
        self.genReportBox.setTitle(_translate("mainWindow", "Generate Destiny Report"))
        self.selectLocation.setItemText(0, _translate("mainWindow", "ABSW"))
        self.selectLocation.setItemText(1, _translate("mainWindow", "Belmont"))
        self.selectLocation.setItemText(2, _translate("mainWindow", "Golden Bear Center"))
        self.selectLocation.setItemText(3, _translate("mainWindow", "San Francisco Center"))
        self.startDateLabel.setText(_translate("mainWindow", "Start Date:"))
        self.browseSaveReportButton.setText(_translate("mainWindow", "Browse"))
        self.selectSaveReportPath.setText(_translate("mainWindow", ""))
        self.locationLabel.setText(_translate("mainWindow", "Location:"))
        self.saveReportPathLabel.setText(_translate("mainWindow", "Save Path:"))
        self.endDateLabel.setText(_translate("mainWindow", "End Date:"))
        self.createSignsBox.setTitle(_translate("mainWindow", "Create Signs"))
        self.singleOutputCheckbox.setText(_translate("mainWindow", "Single File Output"))
        self.useExistingReportBox.setTitle(_translate("mainWindow", "Use existing Destiny Report"))
        self.selectExistingReportPath.setText(_translate("mainWindow", ""))
        self.browseExistingReportButton.setText(_translate("mainWindow", "Browse"))
        self.selectSaveSignsPath.setText(_translate("mainWindow", ""))
        self.browseSaveSignsButton.setText(_translate("mainWindow", "Browse"))
        self.saveSignsPathLabel.setText(_translate("mainWindow", "Save Path:"))
        self.exitButton.setText(_translate("mainWindow", "Exit"))
        self.StartButton.setText(_translate("mainWindow", "Start"))

        self.genReportBox.toggled.connect(self.genReportState)
        self.selectStartDate.dateChanged.connect(self.startDateChanged)
        self.selectEndDate.dateChanged.connect(self.endDateChanged)
        self.selectLocation.currentIndexChanged.connect(self.locationChanged)
        self.browseSaveReportButton.clicked.connect(self.saveReportDirectory)
        self.createSignsBox.toggled.connect(self.createSignsState)
        self.useExistingReportBox.toggled.connect(self.useExistingReportState)
        self.browseExistingReportButton.clicked.connect(self.existingReportPath)
        self.singleOutputCheckbox.toggled.connect(self.singleOutputState)
        self.browseSaveSignsButton.clicked.connect(self.saveSignsPath)
        self.exitButton.clicked.connect(self.exitApp)
        self.StartButton.clicked.connect(self.startApp)

    def genReportState(self):
        global genReport, useExistingReport
        if self.genReportBox.isChecked():
            genReport = True
            useExistingReport = False
            self.useExistingReportBox.setChecked(False)   
        else:   
            genReport = False
            useExistingReport = True
            self.useExistingReportBox.setEnabled(True)
            self.useExistingReportBox.setChecked(True)
        #print('genReport = ' + str(genReport))

    def startDateChanged(self):
        global startDate
        startDate = str(self.selectStartDate.date().toPyDate())
        #print(startDate)

    def endDateChanged(self):
        global endDate
        endDate = str(self.selectEndDate.date().toPyDate())
        #print(endDate)

    def locationChanged(self, i):
        global location
        location = self.selectLocation.currentText()
        #print(location)

    def saveReportDirectory(self):
        global saveReportToPath
        saveReportToPath = QFileDialog.getExistingDirectory(None, 'Save Destiny Report to')
        #print(saveReportToPath)
        self.selectSaveReportPath.setText(saveReportToPath)

    def createSignsState(self):
        global createSigns, useExistingReport
        if self.createSignsBox.isChecked():
            createSigns = True
            if not self.genReportBox.isChecked():
                useExistingReport = True
                self.useExistingReportBox.setChecked(True)
                self.useExistingReportBox.setEnabled(True)
        else:   
            createSigns = False
            useExistingReport = False
            self.useExistingReportBox.setChecked(False)
            self.useExistingReportBox.setEnabled(False)
        #print('createSigns = ' + str(createSigns))

    def useExistingReportState(self):
        global genReport, useExistingReport
        if self.useExistingReportBox.isChecked():
            useExistingReport = True
            genReport = False
            self.genReportBox.setChecked(False)
        else:   
            useExistingReport = False
            genReport = True
            self.genReportBox.setChecked(True)
        #print('useExistingReport = ' + str(useExistingReport))

    def existingReportPath(self):
        global existingReportPath
        existingReportPath, _ = QFileDialog.getOpenFileName(None, "Select SectionScheduleDailySummary.xls", "", "Excel Files (*.xls)")
        #print(existingReportPath)
        self.selectExistingReportPath.setText(existingReportPath)

    def singleOutputState(self):
        global singleOutput
        if self.singleOutputCheckbox.isChecked():
            singleOutput = True
        else:   
            singleOutput = False
        #print('singeOutput = ' + str(singeOutput))

    def saveSignsPath(self):
        global saveSignsDirectory
        saveSignsDirectory = QFileDialog.getExistingDirectory(None, 'Save Signs to')
        #print(saveSignsDirectory)
        self.selectSaveSignsPath.setText(saveSignsDirectory)

    def exitApp(self):        
        reply = QMessageBox.question(None, 'Exit', "Are you sure you want to exit?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)                                      
        if reply == QMessageBox.Yes:
            #print('Yes, exit')
            sys.exit()
        else:
            #print('No, do not exit')
            pass

    def startApp(self):
##        print('genReport: ' + str(genReport))
##        print ('createSigns: ' + str(createSigns))
##        print('startDate: ' + startDate)
##        print('endDate: ' + endDate)
##        print('location: ' + location)
##        print('useExistingReport: ' + str(useExistingReport))
##        print('Save Report to: ' + saveReportToPath)
##        print('existingReportPath: ' + existingReportPath)
##        print('singleOutput: ' + str(singleOutput))
##        print('saveSignsDirectory: ' + saveSignsDirectory)
##        print('')
        
        if genReport == False and createSigns == False:
            QMessageBox.warning(None, 'No settings', "Please fill in the settings first.")
            return
        if genReport:
            if endDate < startDate:
                QMessageBox.warning(None, 'Invalid date range', "Please select a valid date range.")
                return
            elif saveReportToPath == '':
                QMessageBox.warning(None, 'Save location error', "Please select where you want to save the report to.")
                return
            else:
                if os.path.isdir(saveReportToPath):
                    genReportFunction()
                else:  
                    QMessageBox.warning(None, 'Save location error', "The directory you've selected does not exist. Please select where you want to save the report to.")
                    return
        if createSigns:
            chars = set(r'<>?[]:|*')
            #print(any((c in chars) for c in saveSignsDirectory[2:]))
            if saveSignsDirectory == '' or os.path.isdir(saveSignsDirectory) == False:
                QMessageBox.warning(None, 'Save location error', "Please select where you want to save the signs to.")
                return 
            elif genReport:
                if any((c in chars) for c in saveSignsDirectory[2:]):
                    print(r'Filename or path contains: <>?[]:|*')
                    QMessageBox.warning(None, 'Save location error', r'The save path cannot contain any of the following characters: <>?[]:| or *')
                    return
                elif os.path.exists(saveReportToPath + '/SectionScheduleDailySummary.xls') == False:
                    QMessageBox.warning(None, 'Save location error', "The file you've selected does not exist. Please select where you want to save the report to.")
                    return
                else:
                    result = createSignsFunction(saveReportToPath + '/SectionScheduleDailySummary.xls')
            elif useExistingReport and os.path.exists(existingReportPath):
                if any((c in chars) for c in saveSignsDirectory[2:]):
                    print(r'Filename or path contains: <>?[]:|*')
                    QMessageBox.warning(None, 'Save location error', r'The save path cannot contain any of the following characters: <>?[]:| or *')
                    return    
                else:  
                    result = createSignsFunction(existingReportPath)

            elif existingReportPath == '' or os.path.exists(existingReportPath) == False:
                QMessageBox.warning(None, 'No existing report found!!!', "Please select the location of an existing report.")
                return
            if result == 0:
            	QMessageBox.warning(None, 'No classes!!!', "No classes scheduled in date range.")
            else:
            	QMessageBox.warning(None, 'Done', "Done creating signs.")
        
                
def genReportFunction():
    #print('Generating report from ' + startDate + ' to ' + endDate + ' for ' + location + 'which is saved as: ' + saveReportToPath + '/SectionScheduleDailySummary.xls\n'
    # Set Chrome defaults to automate download
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_experimental_option("prefs", {
        "download.default_directory": saveReportToPath,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.endabled": True
        })
        
    # Delete old report if it exists
    if os.path.exists(saveReportToPath + '/SectionScheduleDailySummary.xls'):
        os.remove(saveReportToPath + '/SectionScheduleDailySummary.xls')

    # Download Destiny Report
    chromedriver = os.path.join(current_folder,"chromedriver.exe")
    #browser = webdriver.Chrome("C:\\Python36\\selenium\\webdriver\\chrome\\chromedriver.exe", chrome_options=chrome_options)
    browser = webdriver.Chrome(executable_path = chromedriver, chrome_options=chrome_options)
    browser.get('https://berkeleysv.destinysolutions.com')
    WebDriverWait(browser,60).until(EC.presence_of_element_located((By.ID,"secondMenuCM")))
    browser.get('https://berkeleysv.destinysolutions.com/srs/reporting/sectionScheduleDailySummary.do?method=load')
    startDateElm = browser.find_element_by_css_selector('#startDateRecordString')
    startDateElm.send_keys(startDate)
    endDateElm = browser.find_element_by_css_selector('#endDateRecordString')
    endDateElm.send_keys(endDate)
    campusElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[1]/select')
    campusElm.send_keys(center[location]['campus'])
    buildingElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[2]/select')
    buildingElm.send_keys(center[location]['building'])
    outputTypeElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/select')
    outputTypeElm.send_keys("Output to XLS (Export)")
    generateReportElm = browser.find_element_by_xpath('//*[@id="processReport"]')
    generateReportElm.click()
    while not os.path.exists(saveReportToPath + '/SectionScheduleDailySummary.xls'):
        time.sleep(1)
    browser.close()
    browser.quit()

    xl = DispatchEx("Excel.Application")
    wb = xl.Workbooks.Open(saveReportToPath + '/SectionScheduleDailySummary.xls', ReadOnly=1)

    # Create a new Module and insert the macro code
    mod = wb.VBProject.VBComponents.Add(1)
    mod.CodeModule.AddFromString(excelPrettyReportMacro)

    # Run the new Macro
    xl.Run("Macro2")
    #wb.VBProject.VBComponents.Remove("Macro2")
    try:    
        for i in wb.VBProject.VBComponents:        
            xlmodule = wb.VBProject.VBComponents(i.Name)
            if xlmodule.Type in [1, 2, 3]:            
                wb.VBProject.VBComponents.Remove(xlmodule)

    except Exception as e:
        print(e)

    # Save the workbook and close Excel
    if startDate == endDate:
        wb.SaveAs(os.path.join(os.path.abspath(saveReportToPath), location + " " + startDate + ".xlsx"), FileFormat=51, ConflictResolution=2)
    else:
        wb.SaveAs(os.path.join(os.path.abspath(saveReportToPath), location + " " + startDate + " to " + endDate + ".xlsx"), FileFormat=51, ConflictResolution=2)
    xl.Quit()
    del xl

def createSignsFunction(reportPath):
    print('Creating signs using: ' + reportPath + ' and saving them to: ' + saveSignsDirectory + '\n')
    # Read in courses from Excel
    # 1     B   Date
    # 3     D   Type
    # 4     E   Start Time
    # 6     G   End Time
    # 9     J   Section Number
    # 11    L   Section Title
    # 12    M   Instructor
    # 13    N   Building
    # 15    P   Room
    # 16    Q   Configuration
    # 17    R   Technology
    # 18    S   Section Size
    # 20    U   Notes
    # 22    W   Final Approval
    #excel_file = 'SectionScheduleDailySummary.xls'
    schedule = pd.read_excel(reportPath, header=6, skipfooter=1, usecols=[1,4,6,11,13,15], parse_dates=['Start Time', 'End Time'])
    if schedule.empty:
        return 0
    else:
        location = centerReverse[schedule.iloc[0][4]]['name']
        template = centerReverse[schedule.iloc[0][4]]['template']
        
    startDate = schedule.iloc[0][0].strftime('%Y-%m-%d')
    endDate = schedule.iloc[-1][0].strftime('%Y-%m-%d')
    sortedSchedule = schedule.sort_values(by=['Date','Room','Start Time'])
    sortedSchedule['Date'] = sortedSchedule['Date'].dt.strftime('%B %d, %Y')
    sortedSchedule['Start Time'] = sortedSchedule['Start Time'].dt.strftime('%I:%M%p')
    sortedSchedule['End Time'] = sortedSchedule['End Time'].dt.strftime('%I:%M%p')

##    # Remove Classroom Signs.docx if already exist
##    if os.path.exists('Classroom Signs.docx'):
##        os.remove('Classroom Signs.docx')

    # Initialize variables for 'for loop'
    previousClassroom = ''
    previousDate = None
    dayofweek = ''
    doc = docx.Document(template)
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = 1
    font = doc.styles['Normal'].font
    font.name='Times New Roman'
    font.bold = True
    font.size = Pt(4)

    # Set page orientation, page size, and margins
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE   # Landscape
        section.page_width = 10058400               # Page Width = 11 inches
        section.page_height = 7772400               # Page Height = 8.5 inches
        section.left_margin = 457200                # Left Margin = 0.5 inches
        section.right_margin = 457200               # Right Margin = 0.5 inches
        section.top_margin = 457200                 # Top Margin = 0.5 inches
        section.bottom_margin = 457200              # Bottom Margin = 0.5 inches
        
    # 'For loop' to add Title, Date, Classroom number, and two column table (section title and start/end time) to each page
    for index in range(0, len(sortedSchedule)):
        #print(str(index) + ' ' + sortedSchedule.iloc[index]['Room'] + ' ' + sortedSchedule.iloc[index]['Section Title'])
        newFile = False
        if(index != 0 and not singleOutput and previousDate != sortedSchedule.iloc[index]['Date']):
            if (index != 0 and index != len(sortedSchedule.index)):
                newFile = True
                #addFooter(doc)
                dayofweek = datetime.datetime.strptime(previousDate, '%B %d, %Y').strftime('%A')
                fileDate = datetime.datetime.strptime(previousDate, '%B %d, %Y').strftime('%Y-%m-%d')
                doc.save(saveSignsDirectory + '/' + location + ' ' + fileDate + ' ' + dayofweek + '.docx')
            # Create Classroom Signs, set defaults
            doc = docx.Document(template)
            paragraph_format = doc.styles['Normal'].paragraph_format
            paragraph_format.space_before = 0
            paragraph_format.space_after = 0
            paragraph_format.line_spacing = 1
            font = doc.styles['Normal'].font
            font.name='Times New Roman'
            font.bold = True
            font.size = Pt(4)

            # Set page orientation, page size, and margins
            for section in doc.sections:
                section.orientation = WD_ORIENT.LANDSCAPE   # Landscape
                section.page_width = 10058400               # Page Width = 11 inches
                section.page_height = 7772400               # Page Height = 8.5 inches
                section.left_margin = 457200                # Left Margin = 0.5 inches
                section.right_margin = 457200               # Right Margin = 0.5 inches
                section.top_margin = 457200                 # Top Margin = 0.5 inches
                section.bottom_margin = 457200              # Bottom Margin = 0.5 inches
                
        if (previousClassroom != sortedSchedule.iloc[index]['Room']):
            if (index != 0 and index != len(sortedSchedule.index)):
                if not newFile:
                    #addFooter(doc)       
                    doc.add_page_break()                                # Reached end of page, start new page
                                                        
            para = doc.add_paragraph()
            para.alignment = 1
            run = para.add_run('UC Berkeley Extension')             # Title
            run.alignment = 1
            run.font.size = Pt(72)

            para = doc.add_paragraph()
            para.alignment = 1
            run = para.add_run(sortedSchedule.iloc[index]['Date'].replace(' 0', ' '))  # Date
            run.font.size = Pt(48)

            para = doc.add_paragraph()
            para.alignment = 0
            run = para.add_run(sortedSchedule.iloc[index]['Room'])  # Classroom Number
            run.font.size = Pt(36)

            para = doc.add_paragraph()
            para.alignment = 0
            run = para.add_run('Class:')                            # Class
            run.font.size = Pt(36)

            run = para.add_run('\n')
            run.font.size = Pt(2)

            table = doc.add_table(rows=1, cols=2)                   # Create table to put each course
            table.alignment = 2
            table.allow_autofit = False

            row = table.rows[0]
            row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
            row.cells[1].text = sortedSchedule.iloc[index]['Start Time'].lstrip('0') + ' to ' + sortedSchedule.iloc[index]['End Time'].lstrip('0')
        else:
            row = table.add_row()                                   # add a row if course is in same classroom
            row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
            row.cells[1].text = sortedSchedule.iloc[index]['Start Time'].lstrip('0') + ' to ' + sortedSchedule.iloc[index]['End Time'].lstrip('0')
            
        previousClassroom = sortedSchedule.iloc[index]['Room']
        previousDate = sortedSchedule.iloc[index]['Date']

    # Format table columns
        for cell in table.columns[0].cells:
            cell.width = Inches(6.9)
        for cell in table.columns[1].cells:
            cell.width = Inches(3.1)
    # Change font size of text in table
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(22)
    # End for loop 'for index in range(0, len(sortedSchedule)):'

    #addFooter(doc)  # add footer to last page
    if singleOutput:
        doc.save(saveSignsDirectory + '/' + location + ' ' + startDate + ' to ' + endDate + '.docx')
    else:
        dayofweek = datetime.datetime.strptime(previousDate, '%B %d, %Y').strftime('%A')
        fileDate = datetime.datetime.strptime(previousDate, '%B %d, %Y').strftime('%Y-%m-%d')
        doc.save(saveSignsDirectory + '/' + location + ' ' + fileDate + ' ' + dayofweek + '.docx')

    xl = DispatchEx("Excel.Application")
    wb = xl.Workbooks.Open(os.path.realpath(reportPath), ReadOnly=1)

    # Create a new Module and insert the macro code
    mod = wb.VBProject.VBComponents.Add(1)
    mod.CodeModule.AddFromString(excelClassScheduleMacro)

    # Run the new Macro
    xl.Run("ReformatDestinyReport")
    try:    
        for i in wb.VBProject.VBComponents:        
            xlmodule = wb.VBProject.VBComponents(i.Name)
            if xlmodule.Type in [1, 2, 3]:            
                wb.VBProject.VBComponents.Remove(xlmodule)

    except Exception as e:
        print(e)

    # Save the workbook and close Excel
    wb.SaveAs(os.path.join(os.path.abspath(saveSignsDirectory), location + " Daily Schedule " + startDate + " to " + endDate + ".xlsx"), FileFormat=51, ConflictResolution=2)
    xl.Quit()
    del xl

    ##with pd.option_context('display.max_rows', None, 'display.max_columns', 14):
    ##    print(sortedSchedule)
        
    #if os.path.exists('\SectionScheduleDailySummary.xls'):
    #    os.remove('\SectionScheduleDailySummary.xls')
    return 1

    
def addFooter(doc):                     
    para = doc.add_paragraph('\n'*18)
    para.alignment = 1
    run = para.add_run('Please do not disturb while class is in session\n')
    run.italic = True
    run.font.size = Pt(28)

    run = para.add_run('\n')
    run.font.size = Pt(6)
    
    run = para.add_run('Open Computer Lab\n')
    run.underline = True
    run.font.size = Pt(18)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Open lab Hours: ')
    run.underline = False
    run.font.size = Pt(14)

    run = para.add_run('Monday - Thursday 8:00am to 9:30pm. Friday, Saturday & Sunday 8:00am to 4:30pm\n')
    run.underline = False
    run.bold = False
    run.font.size = Pt(14)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Please Note: Computer Lab will start closing ')
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run('30 minutes')
    run.underline = True
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run(' before closing. Thank you.')
    run.underline = False
    run.italic = True
    run.font.size = Pt(14)
# End def addFooter(doc)      
    
            
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    #print(QtWidgets.QStyleFactory.keys())
    #app.setStyle(QtWidgets.QStyleFactory.create('Fusion'))
    mainWindow = QtWidgets.QWidget()
    ui = Ui_mainWindow()
    ui.setupUi(mainWindow)
    mainWindow.show()
    sys.exit(app.exec_())
