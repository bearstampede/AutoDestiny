#! python3
import os
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from datetime import datetime
import pandas as pd
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches

# Function definition to add footer
def addFooter(doc):                     
    para = doc.add_paragraph('\n'*20)
    para.alignment = 1
    run = para.add_run('Please do not disturb while class is in session\n')
    run.italic = True
    run.font.size = Pt(28)

    run = para.add_run('\n')
    run.font.size = Pt(6)
    
    run = para.add_run('Open Computer Lab\n')
    run.underline = True
    run.font.size = Pt(18)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Open lab Hours: ')
    run.underline = False
    run.font.size = Pt(14)

    run = para.add_run('Monday - Thursday 8:00am to 9:30pm. Friday, Saturday & Sunday 8:00am to 4:30pm\n')
    run.underline = False
    run.bold = False
    run.font.size = Pt(14)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Please Note: Computer Lab will start closing ')
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run('30 minutes')
    run.underline = True
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run(' before closing. Thank you.')
    run.underline = False
    run.italic = True
    run.font.size = Pt(14)
# End def addFooter(doc)  

# Initialize variables
#startDate = '6/20/2018'
#endDate = '6/21/2018'
#downloadPath = r"%userprofile%\Desktop\Test"
#downloadPath = os.path.join(os.getenv('USERPROFILE'), 'Desktop\Test')
downloadPath = os.getcwd()
campus = "Berkeley - CA0001"
building = "UC Berkeley Extension Golden Bear Center, 1995 University Ave. - GBC"
yes = {'yes', 'y', 'Yes', 'YES'}

if(input('Download a new Destiny Report? ') in yes):
    startDate = input('Start Date: ')
    endDate = input('End Date: ')
    # Set Chrome defaults to automate download
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_experimental_option("prefs", {
        "download.default_directory": downloadPath,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.endabled": True
        })
    
    # Delete old report if it exists
    if os.path.exists('SectionScheduleDailySummary.xls'):
        os.remove('SectionScheduleDailySummary.xls')

    # Download Destiny Report
    browser = webdriver.Chrome("C:\\Python36\\selenium\\webdriver\\chrome\\chromedriver.exe", chrome_options=chrome_options)
    browser.get('https://berkeleysv.destinysolutions.com')
    WebDriverWait(browser,60).until(EC.presence_of_element_located((By.ID,"secondMenuCM")))
    browser.get('https://berkeleysv.destinysolutions.com/srs/reporting/sectionScheduleDailySummary.do?method=load')
    startDateElm = browser.find_element_by_css_selector('#startDateRecordString')
    startDateElm.send_keys(startDate)
    endDateElm = browser.find_element_by_css_selector('#endDateRecordString')
    endDateElm.send_keys(endDate)
    campusElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[1]/select')
    campusElm.send_keys(campus)
    buildingElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[2]/select')
    buildingElm.send_keys(building)
    outputTypeElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/select')
    outputTypeElm.send_keys("Output to XLS (Export)")
    generateReportElm = browser.find_element_by_xpath('//*[@id="processReport"]')
    generateReportElm.click()
    while not os.path.exists('SectionScheduleDailySummary.xls'):
        time.sleep(1)
    browser.close()
    browser.quit()

# Read in courses from Excel
# 1     B   Date
# 3     D   Type
# 4     E   Start Time
# 6     G   End Time
# 9     J   Section Number
# 11    L   Section Title
# 12    M   Instructor
# 13    N   Building
# 15    P   Room
# 16    Q   Configuration
# 17    R   Technology
# 18    S   Section Size
# 20    U   Notes
# 22    W   Final Approval
excel_file = 'SectionScheduleDailySummary.xls'
schedule = pd.read_excel(excel_file, header=6, skipfooter=1, usecols=[1,4,6,11,13,15], parse_dates=['Start Time', 'End Time'])
sortedSchedule = schedule.sort_values(by=['Date','Room','Start Time'])
sortedSchedule['Date'] = sortedSchedule['Date'].dt.strftime('%B %d, %Y')
sortedSchedule['Start Time'] = sortedSchedule['Start Time'].dt.strftime('%I:%M%p')
sortedSchedule['End Time'] = sortedSchedule['End Time'].dt.strftime('%I:%M%p')

# Remove Classroom Signs.docx if already exist
if os.path.exists('Classroom Signs.docx'):
    os.remove('Classroom Signs.docx')

# Create Classroom Signs, set defaults
doc = docx.Document('defaultTemplate.docx')
paragraph_format = doc.styles['Normal'].paragraph_format
paragraph_format.space_before = 0
paragraph_format.space_after = 0
paragraph_format.line_spacing = 1
font = doc.styles['Normal'].font
font.name='Times New Roman'
font.bold = True
font.size = Pt(4)

# Set page orientation, page size, and margins
for section in doc.sections:
    section.orientation = WD_ORIENT.LANDSCAPE   # Landscape
    section.page_width = 10058400               # Page Width = 11 inches
    section.page_height = 7772400               # Page Height = 8.5 inches
    section.left_margin = 457200                # Left Margin = 0.5 inches
    section.right_margin = 457200               # Right Margin = 0.5 inches
    section.top_margin = 457200                 # Top Margin = 0.5 inches
    section.bottom_margin = 457200              # Bottom Margin = 0.5 inches

# Initialize variables for 'for loop'
previousClassroom = ''
previousDate = ''

# 'For loop' to add Title, Date, Classroom number, and two column table (section title and start/end time) to each page
for index in range(0, len(sortedSchedule)):
    #print(str(index) + ' ' + sortedSchedule.iloc[index]['Room'] + ' ' + sortedSchedule.iloc[index]['Section Title'])
    if (previousClassroom != sortedSchedule.iloc[index]['Room'] or previousDate != sortedSchedule.iloc[index]['Date']):
        if (index != 0 and index != len(sortedSchedule.index)):
            addFooter(doc)       
            doc.add_page_break()                                # Reached end of page, start new page
                                                    
        para = doc.add_paragraph()
        para.alignment = 1
        run = para.add_run('UC Berkeley Extension')             # Title
        run.alignment = 1
        run.font.size = Pt(72)

        para = doc.add_paragraph()
        para.alignment = 1
        run = para.add_run(sortedSchedule.iloc[index]['Date'])  # Date
        run.font.size = Pt(48)

        para = doc.add_paragraph()
        para.alignment = 0
        run = para.add_run(sortedSchedule.iloc[index]['Room'])  # Classroom Number
        run.font.size = Pt(36)

        para = doc.add_paragraph()
        para.alignment = 0
        run = para.add_run('Class:')                            # Class
        run.font.size = Pt(36)

        table = doc.add_table(rows=1, cols=2)                   # Create table to put each course
        table.alignment = 1
        table.allow_autofit = False

        row = table.rows[0]
        row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
        row.cells[1].text = sortedSchedule.iloc[index]['Start Time'].lstrip('0') + ' to ' + sortedSchedule.iloc[index]['End Time'].lstrip('0')
    else:
        row = table.add_row()                                   # add a row if course is in same classroom
        row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
        row.cells[1].text = sortedSchedule.iloc[index]['Start Time'].lstrip('0') + ' to ' + sortedSchedule.iloc[index]['End Time'].lstrip('0')
        
    previousClassroom = sortedSchedule.iloc[index]['Room']
    previousDate = sortedSchedule.iloc[index]['Date']

# Format table columns
    for cell in table.columns[0].cells:
        cell.width = Inches(7)
    for cell in table.columns[1].cells:
        cell.width = Inches(3)
# Change font size of text in table
    for row in table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.size = Pt(22)
# End for loop 'for index in range(0, len(sortedSchedule)):'

addFooter(doc)  # add footer to last page

doc.save('Classroom Signs.docx')

##with pd.option_context('display.max_rows', None, 'display.max_columns', 14):
##    print(sortedSchedule)
    
#if os.path.exists('\SectionScheduleDailySummary.xls'):
#    os.remove('\SectionScheduleDailySummary.xls')



