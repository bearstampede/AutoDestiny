#! python3
import sys, inspect
import datetime
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtWidgets import QFileDialog, QApplication, QWidget, QPushButton, QMessageBox
import os
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import time
from datetime import datetime
import pandas as pd
import docx
from docx.enum.section import WD_ORIENT
from docx.shared import Pt, Inches

genReport = False
startDate = '2018-01-01'
endDate = '2018-01-01'
location = 'ABSW'
createSigns = False
useExistingReport = False
saveReportToPath = ''
existingReportPath, _ = ('','')
singleOutput = False
saveSignsDirectory = ''
center = {'ABSW': {'campus': 'Berkeley - CA0001', 'building': 'UC Berkeley Extension American Baptist Seminary of the West, 2515 Hillegass Ave. - ABSW'},
          'Belmont': {'campus': 'Belmont - CA0004', 'building': 'UC Berkeley Extension Belmont Center, 1301 Shoreway Rd., Ste. 400 - BEL'},
          'Golden Bear Center': {'campus': 'Berkeley - CA0001', 'building': 'UC Berkeley Extension Golden Bear Center, 1995 University Ave. - GBC'},
          'San Francisco Center': {'campus': 'San Francisco - CA0003', 'building': 'San Francisco Campus, 160 Spear St. - SFCAMPUS'}
            }

current_folder = os.getcwd()			
def init_driver():
    chromedriver = os.path.join(current_folder, "chromedriver.exe")
    # via this way, you explicitly let Chrome know where to find 
    # the webdriver.
    driver = webdriver.Chrome(current_folder, chrome_options = chrome_options) 
    return driver
			
# Main Window for GUI
class Ui_mainWindow(object): 
    def setupUi(self, mainWindow):
        global startDate
        global endDate
        mainWindow.setObjectName("mainWindow")
        mainWindow.setWindowModality(QtCore.Qt.NonModal)
        mainWindow.setEnabled(True)
        mainWindow.resize(430, 380)
        sizePolicy = QtWidgets.QSizePolicy(QtWidgets.QSizePolicy.Fixed, QtWidgets.QSizePolicy.Fixed)
        sizePolicy.setHorizontalStretch(0)
        sizePolicy.setVerticalStretch(0)
        sizePolicy.setHeightForWidth(mainWindow.sizePolicy().hasHeightForWidth())
        mainWindow.setSizePolicy(sizePolicy)
        mainWindow.setMinimumSize(QtCore.QSize(430, 380))
        mainWindow.setMaximumSize(QtCore.QSize(430, 380))
        mainWindow.setBaseSize(QtCore.QSize(430, 325))
        self.gridLayoutWidget = QtWidgets.QWidget(mainWindow)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(10, 10, 412, 371))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.widgetVerticalLayout = QtWidgets.QVBoxLayout(self.gridLayoutWidget)
        self.widgetVerticalLayout.setSizeConstraint(QtWidgets.QLayout.SetDefaultConstraint)
        self.widgetVerticalLayout.setContentsMargins(0, 0, 0, 0)
        self.widgetVerticalLayout.setObjectName("widgetVerticalLayout")
        self.genReportBox = QtWidgets.QGroupBox(self.gridLayoutWidget)
        self.genReportBox.setMinimumSize(QtCore.QSize(410, 120))
        self.genReportBox.setMaximumSize(QtCore.QSize(410, 120))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(14)
        font.setUnderline(True)
        self.genReportBox.setFont(font)
        self.genReportBox.setCheckable(True)
        self.genReportBox.setChecked(False)
        self.genReportBox.setObjectName("genReportBox")
        self.selectLocation = QtWidgets.QComboBox(self.genReportBox)
        self.selectLocation.setGeometry(QtCore.QRect(90, 60, 171, 22))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectLocation.setFont(font)
        self.selectLocation.setEditable(False)
        self.selectLocation.setObjectName("selectLocation")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.selectLocation.addItem("")
        self.startDateLabel = QtWidgets.QLabel(self.genReportBox)
        self.startDateLabel.setGeometry(QtCore.QRect(20, 30, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.startDateLabel.setFont(font)
        self.startDateLabel.setObjectName("startDateLabel")
        self.browseSaveReportButton = QtWidgets.QToolButton(self.genReportBox)
        self.browseSaveReportButton.setGeometry(QtCore.QRect(340, 90, 61, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setUnderline(False)
        self.browseSaveReportButton.setFont(font)
        self.browseSaveReportButton.setCheckable(False)
        self.browseSaveReportButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseSaveReportButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseSaveReportButton.setObjectName("browseSaveReportButton")
        self.selectSaveReportPath = QtWidgets.QLineEdit(self.genReportBox)
        self.selectSaveReportPath.setGeometry(QtCore.QRect(100, 90, 241, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectSaveReportPath.setFont(font)
        self.selectSaveReportPath.setReadOnly(False)
        self.selectSaveReportPath.setObjectName("selectSaveReportPath")
        self.selectStartDate = QtWidgets.QDateEdit(self.genReportBox)
        self.selectStartDate.setGeometry(QtCore.QRect(100, 30, 111, 22))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setBold(False)
        font.setUnderline(False)
        font.setWeight(50)
        self.selectStartDate.setFont(font)
        self.selectStartDate.setFrame(True)
        self.selectStartDate.setProperty("showGroupSeparator", False)
        self.selectStartDate.setCalendarPopup(True)
        self.selectStartDate.setDate(QtCore.QDate.currentDate())
        startDate = str(QtCore.QDate.currentDate().toPyDate())
        self.selectStartDate.setObjectName("selectStartDate")
        self.locationLabel = QtWidgets.QLabel(self.genReportBox)
        self.locationLabel.setGeometry(QtCore.QRect(20, 60, 71, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.locationLabel.setFont(font)
        self.locationLabel.setObjectName("locationLabel")
        self.saveReportPathLabel = QtWidgets.QLabel(self.genReportBox)
        self.saveReportPathLabel.setGeometry(QtCore.QRect(20, 90, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.saveReportPathLabel.setFont(font)
        self.saveReportPathLabel.setObjectName("saveReportPathLabel")
        self.endDateLabel = QtWidgets.QLabel(self.genReportBox)
        self.endDateLabel.setGeometry(QtCore.QRect(220, 30, 71, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setBold(False)
        font.setUnderline(False)
        font.setWeight(50)
        self.endDateLabel.setFont(font)
        self.endDateLabel.setObjectName("endDateLabel")
        self.selectEndDate = QtWidgets.QDateEdit(self.genReportBox)
        self.selectEndDate.setGeometry(QtCore.QRect(290, 30, 110, 22))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectEndDate.setFont(font)
        self.selectEndDate.setFocusPolicy(QtCore.Qt.WheelFocus)
        self.selectEndDate.setCalendarPopup(True)
        self.selectEndDate.setDate(QtCore.QDate.currentDate())
        endDate = str(QtCore.QDate.currentDate().toPyDate())
        self.selectEndDate.setObjectName("selectEndDate")
        self.selectLocation.raise_()
        self.startDateLabel.raise_()
        self.browseSaveReportButton.raise_()
        self.selectSaveReportPath.raise_()
        self.selectStartDate.raise_()
        self.locationLabel.raise_()
        self.saveReportPathLabel.raise_()
        self.endDateLabel.raise_()
        self.selectEndDate.raise_()
        self.widgetVerticalLayout.addWidget(self.genReportBox)
        self.createSignsBox = QtWidgets.QGroupBox(self.gridLayoutWidget)
        self.createSignsBox.setEnabled(True)
        self.createSignsBox.setMinimumSize(QtCore.QSize(150, 160))
        self.createSignsBox.setMaximumSize(QtCore.QSize(410, 150))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(14)
        font.setUnderline(True)
        self.createSignsBox.setFont(font)
        self.createSignsBox.setCheckable(True)
        self.createSignsBox.setChecked(False)
        self.createSignsBox.setObjectName("createSignsBox")
        self.singleOutputCheckbox = QtWidgets.QCheckBox(self.createSignsBox)
        self.singleOutputCheckbox.setGeometry(QtCore.QRect(20, 100, 151, 21))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.singleOutputCheckbox.setFont(font)
        self.singleOutputCheckbox.setChecked(False)
        self.singleOutputCheckbox.setObjectName("singleOutputCheckbox")
        self.useExistingReportBox = QtWidgets.QGroupBox(self.createSignsBox)
        self.useExistingReportBox.setGeometry(QtCore.QRect(10, 30, 390, 60))
        self.useExistingReportBox.setMaximumSize(QtCore.QSize(390, 60))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.useExistingReportBox.setFont(font)
        self.useExistingReportBox.setCheckable(True)
        self.useExistingReportBox.setChecked(False)
        self.useExistingReportBox.setObjectName("useExistingReportBox")
        self.selectExistingReportPath = QtWidgets.QLineEdit(self.useExistingReportBox)
        self.selectExistingReportPath.setGeometry(QtCore.QRect(10, 30, 311, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        self.selectExistingReportPath.setFont(font)
        self.selectExistingReportPath.setReadOnly(False)
        self.selectExistingReportPath.setObjectName("selectExistingReportPath")
        self.browseExistingReportButton = QtWidgets.QToolButton(self.useExistingReportBox)
        self.browseExistingReportButton.setGeometry(QtCore.QRect(320, 30, 61, 19))
        font = QtGui.QFont()
        font.setPointSize(10)
        self.browseExistingReportButton.setFont(font)
        self.browseExistingReportButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseExistingReportButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseExistingReportButton.setAutoRaise(False)
        self.browseExistingReportButton.setArrowType(QtCore.Qt.NoArrow)
        self.browseExistingReportButton.setObjectName("browseExistingReportButton")
        self.selectExistingReportPath.raise_()
        self.browseExistingReportButton.raise_()
        self.genReportBox.raise_()
        self.genReportBox.raise_()
        self.selectSaveSignsPath = QtWidgets.QLineEdit(self.createSignsBox)
        self.selectSaveSignsPath.setGeometry(QtCore.QRect(100, 130, 231, 20))
        font = QtGui.QFont()
        font.setPointSize(12)
        font.setUnderline(False)
        self.selectSaveSignsPath.setFont(font)
        self.selectSaveSignsPath.setReadOnly(False)
        self.selectSaveSignsPath.setObjectName("selectSaveSignsPath")
        self.browseSaveSignsButton = QtWidgets.QToolButton(self.createSignsBox)
        self.browseSaveSignsButton.setGeometry(QtCore.QRect(330, 130, 61, 20))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setUnderline(False)
        self.browseSaveSignsButton.setFont(font)
        self.browseSaveSignsButton.setCheckable(False)
        self.browseSaveSignsButton.setPopupMode(QtWidgets.QToolButton.InstantPopup)
        self.browseSaveSignsButton.setToolButtonStyle(QtCore.Qt.ToolButtonTextOnly)
        self.browseSaveSignsButton.setObjectName("browseSaveSignsButton")
        self.saveSignsPathLabel = QtWidgets.QLabel(self.createSignsBox)
        self.saveSignsPathLabel.setGeometry(QtCore.QRect(20, 130, 81, 16))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        font.setUnderline(False)
        self.saveSignsPathLabel.setFont(font)
        self.saveSignsPathLabel.setObjectName("saveSignsPathLabel")
        self.widgetVerticalLayout.addWidget(self.createSignsBox)
        self.startExitLayout = QtWidgets.QGridLayout()
        self.startExitLayout.setObjectName("startExitLayout")
        self.exitButton = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.exitButton.setMinimumSize(QtCore.QSize(0, 50))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        self.exitButton.setFont(font)
        self.exitButton.setObjectName("exitButton")
        self.startExitLayout.addWidget(self.exitButton, 0, 1, 1, 1)
        self.StartButton = QtWidgets.QPushButton(self.gridLayoutWidget)
        self.StartButton.setMinimumSize(QtCore.QSize(0, 50))
        font = QtGui.QFont()
        font.setFamily("Tahoma")
        font.setPointSize(12)
        self.StartButton.setFont(font)
        self.StartButton.setObjectName("StartButton")
        self.startExitLayout.addWidget(self.StartButton, 0, 0, 1, 1)
        self.widgetVerticalLayout.addLayout(self.startExitLayout)

        self.retranslateUi(mainWindow)
        QtCore.QMetaObject.connectSlotsByName(mainWindow)

    def retranslateUi(self, mainWindow):
        _translate = QtCore.QCoreApplication.translate
        mainWindow.setWindowTitle(_translate("mainWindow", "Auto Schedule"))
        self.genReportBox.setTitle(_translate("mainWindow", "Generate Destiny Report"))
        self.selectLocation.setItemText(0, _translate("mainWindow", "ABSW"))
        self.selectLocation.setItemText(1, _translate("mainWindow", "Belmont"))
        self.selectLocation.setItemText(2, _translate("mainWindow", "Golden Bear Center"))
        self.selectLocation.setItemText(3, _translate("mainWindow", "San Francisco Center"))
        self.startDateLabel.setText(_translate("mainWindow", "Start Date:"))
        self.browseSaveReportButton.setText(_translate("mainWindow", "Browse"))
        self.selectSaveReportPath.setText(_translate("mainWindow", ""))
        self.locationLabel.setText(_translate("mainWindow", "Location:"))
        self.saveReportPathLabel.setText(_translate("mainWindow", "Save Path:"))
        self.endDateLabel.setText(_translate("mainWindow", "End Date:"))
        self.createSignsBox.setTitle(_translate("mainWindow", "Create Signs"))
        self.singleOutputCheckbox.setText(_translate("mainWindow", "Single File Output"))
        self.useExistingReportBox.setTitle(_translate("mainWindow", "Use existing Destiny Report"))
        self.selectExistingReportPath.setText(_translate("mainWindow", ""))
        self.browseExistingReportButton.setText(_translate("mainWindow", "Browse"))
        self.selectSaveSignsPath.setText(_translate("mainWindow", ""))
        self.browseSaveSignsButton.setText(_translate("mainWindow", "Browse"))
        self.saveSignsPathLabel.setText(_translate("mainWindow", "Save Path:"))
        self.exitButton.setText(_translate("mainWindow", "Exit"))
        self.StartButton.setText(_translate("mainWindow", "Start"))

        self.genReportBox.toggled.connect(self.genReportState)
        self.selectStartDate.dateChanged.connect(self.startDateChanged)
        self.selectEndDate.dateChanged.connect(self.endDateChanged)
        self.selectLocation.currentIndexChanged.connect(self.locationChanged)
        self.browseSaveReportButton.clicked.connect(self.saveReportDirectory)
        self.createSignsBox.toggled.connect(self.createSignsState)
        self.useExistingReportBox.toggled.connect(self.useExistingReportState)
        self.browseExistingReportButton.clicked.connect(self.existingReportPath)
        self.singleOutputCheckbox.toggled.connect(self.singleOutputState)
        self.browseSaveSignsButton.clicked.connect(self.saveSignsPath)
        self.exitButton.clicked.connect(self.exitApp)
        self.StartButton.clicked.connect(self.startApp)

    def genReportState(self):
        global genReport, useExistingReport
        if self.genReportBox.isChecked():
            genReport = True
            useExistingReport = False
            self.useExistingReportBox.setChecked(False)   
        else:   
            genReport = False
            useExistingReport = True
            self.useExistingReportBox.setEnabled(True)
            self.useExistingReportBox.setChecked(True)
        #print('genReport = ' + str(genReport))

    def startDateChanged(self):
        global startDate
        startDate = str(self.selectStartDate.date().toPyDate())
        #print(startDate)

    def endDateChanged(self):
        global endDate
        endDate = str(self.selectEndDate.date().toPyDate())
        #print(endDate)

    def locationChanged(self, i):
        global location
        location = self.selectLocation.currentText()
        #print(location)

    def saveReportDirectory(self):
        global saveReportToPath
        saveReportToPath = QFileDialog.getExistingDirectory(None, 'Save Destiny Report to')
        #print(saveReportToPath)
        self.selectSaveReportPath.setText(saveReportToPath)

    def createSignsState(self):
        global createSigns, useExistingReport
        if self.createSignsBox.isChecked():
            createSigns = True
            if not self.genReportBox.isChecked():
                useExistingReport = True
                self.useExistingReportBox.setChecked(True)
                self.useExistingReportBox.setEnabled(True)
        else:   
            createSigns = False
            useExistingReport = False
            self.useExistingReportBox.setChecked(False)
            self.useExistingReportBox.setEnabled(False)
        #print('createSigns = ' + str(createSigns))

    def useExistingReportState(self):
        global genReport, useExistingReport
        if self.useExistingReportBox.isChecked():
            useExistingReport = True
            genReport = False
            self.genReportBox.setChecked(False)
        else:   
            useExistingReport = False
            genReport = True
            self.genReportBox.setChecked(True)
        #print('useExistingReport = ' + str(useExistingReport))

    def existingReportPath(self):
        global existingReportPath
        existingReportPath, _ = QFileDialog.getOpenFileName(None, "Select SectionScheduleDailySummary.xls", "", "Excel Files (*.xls)")
        #print(existingReportPath)
        self.selectExistingReportPath.setText(existingReportPath)

    def singleOutputState(self):
        global singleOutput
        if self.singleOutputCheckbox.isChecked():
            singleOutput = True
        else:   
            singleOutput = False
        #print('singeOutput = ' + str(singeOutput))

    def saveSignsPath(self):
        global saveSignsDirectory
        saveSignsDirectory = QFileDialog.getExistingDirectory(None, 'Save Signs to')
        #print(saveSignsDirectory)
        self.selectSaveSignsPath.setText(saveSignsDirectory)

    def exitApp(self):        
        reply = QMessageBox.question(None, 'Exit', "Are you sure you want to exit?", QMessageBox.Yes | QMessageBox.No, QMessageBox.No)                                      
        if reply == QMessageBox.Yes:
            #print('Yes, exit')
            sys.exit()
        else:
            #print('No, do not exit')
            pass

    def startApp(self):
##        print('genReport: ' + str(genReport))
##        print ('createSigns: ' + str(createSigns))
##        print('startDate: ' + startDate)
##        print('endDate: ' + endDate)
##        print('location: ' + location)
##        print('useExistingReport: ' + str(useExistingReport))
##        print('Save Report to: ' + saveReportToPath)
##        print('existingReportPath: ' + existingReportPath)
##        print('singleOutput: ' + str(singleOutput))
##        print('saveSignsDirectory: ' + saveSignsDirectory)
##        print('')
        
        if genReport == False and createSigns == False:
            QMessageBox.warning(None, 'No settings', "Please fill in the settings first.")
        if genReport:
            if endDate < startDate:
                QMessageBox.warning(None, 'Invalid date range', "Please select a valid date range.")
            elif saveReportToPath == '':
                QMessageBox.warning(None, 'Save location error', "Please select where you want to save the report to.")
            else:
                genReportFunction()
        if createSigns:
            if saveSignsDirectory == '':
                QMessageBox.warning(None, 'Save location error', "Please select where you want to save the signs to.")
            elif genReport:
                createSignsFunction(saveReportToPath + '/SectionScheduleDailySummary.xls')
            elif useExistingReport and existingReportPath:
                createSignsFunction(existingReportPath)
            elif existingReportPath == '':
                QMessageBox.warning(None, 'No existing report found!!!', "Please select the location of an existing report.")
                
def genReportFunction():
    #print('Generating report from ' + startDate + ' to ' + endDate + ' for ' + location + 'which is saved as: ' + saveReportToPath + '/SectionScheduleDailySummary.xls\n'
    # Set Chrome defaults to automate download
    chrome_options = webdriver.ChromeOptions()
    chrome_options.add_experimental_option("prefs", {
        "download.default_directory": saveReportToPath,
        "download.prompt_for_download": False,
        "download.directory_upgrade": True,
        "safebrowsing.enabled": True
        })
        
    # Delete old report if it exists
    if os.path.exists(saveReportToPath + '/SectionScheduleDailySummary.xls'):
        os.remove(saveReportToPath + '/SectionScheduleDailySummary.xls')

    # Download Destiny Report
    browser = webdriver.Chrome(current_folder, chrome_options = chrome_options) 
    browser.get('https://berkeleysv.destinysolutions.com')
    WebDriverWait(browser,60).until(EC.presence_of_element_located((By.ID,"secondMenuCM")))
    browser.get('https://berkeleysv.destinysolutions.com/srs/reporting/sectionScheduleDailySummary.do?method=load')
    startDateElm = browser.find_element_by_css_selector('#startDateRecordString')
    startDateElm.send_keys(startDate)
    endDateElm = browser.find_element_by_css_selector('#endDateRecordString')
    endDateElm.send_keys(endDate)
    campusElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[1]/select')
    campusElm.send_keys(center[location]['campus'])
    buildingElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/table/tbody/tr[3]/td/table/tbody/tr[2]/td[2]/select')
    buildingElm.send_keys(center[location]['building'])
    outputTypeElm = browser.find_element_by_xpath('//*[@id="content01"]/table/tbody/tr[2]/td/table/tbody/tr/td/select')
    outputTypeElm.send_keys("Output to XLS (Export)")
    generateReportElm = browser.find_element_by_xpath('//*[@id="processReport"]')
    generateReportElm.click()
    while not os.path.exists(saveReportToPath + '/SectionScheduleDailySummary.xls'):
        time.sleep(1)
    browser.close()
    browser.quit()


def createSignsFunction(reportPath):
    print('Creating signs using: ' + reportPath + ' and saving them to: ' + saveSignsDirectory + '\n')
    # Read in courses from Excel
    # 1     B   Date
    # 3     D   Type
    # 4     E   Start Time
    # 6     G   End Time
    # 9     J   Section Number
    # 11    L   Section Title
    # 12    M   Instructor
    # 13    N   Building
    # 15    P   Room
    # 16    Q   Configuration
    # 17    R   Technology
    # 18    S   Section Size
    # 20    U   Notes
    # 22    W   Final Approval
    #excel_file = 'SectionScheduleDailySummary.xls'
    schedule = pd.read_excel(reportPath, header=6, skipfooter=1, usecols=[1,4,6,11,15], parse_dates=['Start Time', 'End Time'])
    sortedSchedule = schedule.sort_values(by=['Date','Room','Start Time'])
    sortedSchedule['Date'] = sortedSchedule['Date'].dt.strftime('%B %d, %Y')
    sortedSchedule['Start Time'] = sortedSchedule['Start Time'].dt.strftime('%I:%M%p')
    sortedSchedule['End Time'] = sortedSchedule['End Time'].dt.strftime('%I:%M%p')

##    # Remove Classroom Signs.docx if already exist
##    if os.path.exists('Classroom Signs.docx'):
##        os.remove('Classroom Signs.docx')

    # Create Classroom Signs, set defaults
    doc = docx.Document('defaultTemplate.docx')
    paragraph_format = doc.styles['Normal'].paragraph_format
    paragraph_format.space_before = 0
    paragraph_format.space_after = 0
    paragraph_format.line_spacing = 1
    font = doc.styles['Normal'].font
    font.name='Times New Roman'
    font.bold = True
    font.size = Pt(4)

    # Set page orientation, page size, and margins
    for section in doc.sections:
        section.orientation = WD_ORIENT.LANDSCAPE   # Landscape
        section.page_width = 10058400               # Page Width = 11 inches
        section.page_height = 7772400               # Page Height = 8.5 inches
        section.left_margin = 457200                # Left Margin = 0.5 inches
        section.right_margin = 457200               # Right Margin = 0.5 inches
        section.top_margin = 457200                 # Top Margin = 0.5 inches
        section.bottom_margin = 457200              # Bottom Margin = 0.5 inches

    # Initialize variables for 'for loop'
    previousClassroom = ''
    previousDate = ''

    # 'For loop' to add Title, Date, Classroom number, and two column table (section title and start/end time) to each page
    for index in range(0, len(sortedSchedule)):
        #print(str(index) + ' ' + sortedSchedule.iloc[index]['Room'] + ' ' + sortedSchedule.iloc[index]['Section Title'])
        if (previousClassroom != sortedSchedule.iloc[index]['Room'] or previousDate != sortedSchedule.iloc[index]['Date']):
            if (index != 0 and index != len(sortedSchedule.index)):
                addFooter(doc)       
                doc.add_page_break()                                # Reached end of page, start new page
                                                        
            para = doc.add_paragraph()
            para.alignment = 1
            run = para.add_run('UC Berkeley Extension')             # Title
            run.alignment = 1
            run.font.size = Pt(72)

            para = doc.add_paragraph()
            para.alignment = 1
            run = para.add_run(sortedSchedule.iloc[index]['Date'])  # Date
            run.font.size = Pt(48)

            para = doc.add_paragraph()
            para.alignment = 0
            run = para.add_run(sortedSchedule.iloc[index]['Room'])  # Classroom Number
            run.font.size = Pt(36)

            para = doc.add_paragraph()
            para.alignment = 0
            run = para.add_run('Class:')                            # Class
            run.font.size = Pt(36)

            run = para.add_run('\n'*2)
            run.font.size = Pt(6)

            table = doc.add_table(rows=1, cols=2)                   # Create table to put each course
            table.alignment = 2
            table.allow_autofit = False

            row = table.rows[0]
            row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
            row.cells[1].text = sortedSchedule.iloc[index]['Start Time'] + ' to ' + sortedSchedule.iloc[index]['End Time']
        else:
            row = table.add_row()                                   # add a row if course is in same classroom
            row.cells[0].text = sortedSchedule.iloc[index]['Section Title'] + '\n'
            row.cells[1].text = sortedSchedule.iloc[index]['Start Time'] + ' to ' + sortedSchedule.iloc[index]['End Time']
            
        previousClassroom = sortedSchedule.iloc[index]['Room']
        previousDate = sortedSchedule.iloc[index]['Date']

    # Format table columns
        for cell in table.columns[0].cells:
            cell.width = Inches(7)
        for cell in table.columns[1].cells:
            cell.width = Inches(3)
    # Change font size of text in table
        for row in table.rows:
            for cell in row.cells:
                paragraphs = cell.paragraphs
                for paragraph in paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.size = Pt(22)
    # End for loop 'for index in range(0, len(sortedSchedule)):'

    addFooter(doc)  # add footer to last page

    doc.save(saveSignsDirectory + '/Classroom Signs.docx')

    ##with pd.option_context('display.max_rows', None, 'display.max_columns', 14):
    ##    print(sortedSchedule)
        
    #if os.path.exists('\SectionScheduleDailySummary.xls'):
    #    os.remove('\SectionScheduleDailySummary.xls')

    
def addFooter(doc):                     
    para = doc.add_paragraph('\n'*20)
    para.alignment = 1
    run = para.add_run('Please do not disturb while class is in session\n')
    run.italic = True
    run.font.size = Pt(28)

    run = para.add_run('\n')
    run.font.size = Pt(6)
    
    run = para.add_run('Open Computer Lab\n')
    run.underline = True
    run.font.size = Pt(18)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Open lab Hours: ')
    run.underline = False
    run.font.size = Pt(14)

    run = para.add_run('Monday - Thursday 8:00am to 9:30pm. Friday, Saturday & Sunday 8:00am to 4:30pm\n')
    run.underline = False
    run.bold = False
    run.font.size = Pt(14)

    run = para.add_run('\n')
    run.font.size = Pt(6)

    run = para.add_run('Please Note: Computer Lab will start closing ')
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run('30 minutes')
    run.underline = True
    run.italic = True
    run.font.size = Pt(14)
    run = para.add_run(' before closing. Thank you.')
    run.underline = False
    run.italic = True
    run.font.size = Pt(14)
# End def addFooter(doc)      
    
            
if __name__ == "__main__":
    app = QtWidgets.QApplication(sys.argv)
    mainWindow = QtWidgets.QWidget()
    ui = Ui_mainWindow()
    ui.setupUi(mainWindow)
    mainWindow.show()
    sys.exit(app.exec_())

